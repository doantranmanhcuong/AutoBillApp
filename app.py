import streamlit as st
import os
import pandas as pd
import zipfile
import io
import tempfile 
import re
from PIL import Image
from docx import Document
import openpyxl
import fitz 

from core.config import API_KEY, TEMPLATE_DIR
from core.utils.helpers import doc_so_tien_vn
from core.ai_extractor import AIExtractor
from core.document_builder import DocumentBuilder
from ui.components import display_zalo_message

st.set_page_config(page_title="Hệ thống ERP AI", page_icon="🧾", layout="wide")
available_templates = [f for f in os.listdir(TEMPLATE_DIR) if f.endswith(('.xlsx', '.docx'))]

st.title("🧾 HỆ THỐNG KẾ TOÁN AI - TẠO BỘ HỒ SƠ TỰ ĐỘNG")
st.markdown("---")

col_top1, col_top2 = st.columns(2)
with col_top1:
    uploaded_file = st.file_uploader("📥 1. Tải lên Hóa đơn / Báo giá", type=["png", "jpg", "jpeg", "xlsx", "docx", "pdf"])
with col_top2:
    if available_templates:
        selected_templates = st.multiselect("📂 2. Chọn Các Biểu Mẫu Muốn Xuất:", available_templates)
    else:
        st.error("⚠️ Chưa có file mẫu nào.")
        selected_templates = []

st.divider()

col_left, col_right = st.columns([5, 5], gap="large")

# ==========================================
# HÀM QUÉT TAG TRIỆT ĐỂ (DEEP SCAN SIÊU MẠNH)
# ==========================================
def exhaustive_extract_tags(file_path):
    tags = set()
    # Biểu thức này quét bắt được mọi loại: {tag}, {{tag}}, {{{ tag }}} kể cả có dấu cách thừa
    pattern = re.compile(r'\{+\s*([a-zA-Z0-9_]+)\s*\}+')
    try:
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            for p in doc.paragraphs:
                tags.update(pattern.findall(p.text))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            tags.update(pattern.findall(p.text))
        elif file_path.endswith('.xlsx'):
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    for cell in row:
                        if isinstance(cell, str):
                            tags.update(pattern.findall(cell))
    except Exception:
        pass
    # Xóa khoảng trắng và chuyển chữ thường để chuẩn hóa 100%
    return {t.strip().lower() for t in tags}

with col_left:
    st.markdown("### 🔍 TÀI LIỆU GỐC")
    if uploaded_file:
        file_ext = uploaded_file.name.split('.')[-1].lower()
        
        if 'uploaded_filename' not in st.session_state or st.session_state['uploaded_filename'] != uploaded_file.name:
            if 'temp_path' in st.session_state and os.path.exists(st.session_state['temp_path']):
                try: os.remove(st.session_state['temp_path'])
                except Exception: pass
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as tmp_file:
                tmp_file.write(uploaded_file.getbuffer())
                st.session_state['temp_path'] = tmp_file.name
                st.session_state['uploaded_filename'] = uploaded_file.name

        temp_input = st.session_state['temp_path']
            
        with st.container(border=True):
            if file_ext in ['png', 'jpg', 'jpeg']: 
                st.image(Image.open(temp_input), use_container_width=True)
            elif file_ext == 'pdf':
                try:
                    doc = fitz.open(temp_input)
                    st.info(f"📄 Tài liệu PDF có tổng cộng {len(doc)} trang.")
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        pix = page.get_pixmap(dpi=150)
                        img_path = f"temp_page_{page_num}.png"
                        pix.save(img_path)
                        st.image(Image.open(img_path), caption=f"Trang {page_num + 1}", use_container_width=True)
                        if os.path.exists(img_path): os.remove(img_path)
                except Exception as e:
                    st.error(f"Không thể hiển thị bản xem trước PDF: {str(e)}")
                    with open(temp_input, "rb") as f:
                        st.download_button("📥 Tải xuống file PDF gốc để xem", f, file_name=uploaded_file.name)
            elif file_ext == 'xlsx':
                for sheet, df in pd.read_excel(temp_input, sheet_name=None).items(): 
                    st.dataframe(df, height=800, use_container_width=True)
            elif file_ext == 'docx':
                text = "\n".join([p.text for p in Document(temp_input).paragraphs])
                st.text_area("Nội dung Word:", text, height=850)

with col_right:
    st.markdown("### ✍️ KIỂM DUYỆT THÔNG TIN")
    if uploaded_file and selected_templates:
        
        template_tags_map = {}
        all_dynamic_tags = set()
        
        # Các thẻ hệ thống thuộc về Bảng Kê & Tổng tiền -> Không hiện ra ở các ô riêng
        TABLE_TAGS = [
            "stt", "ten_hang_hoa", "don_vi_tinh", "so_luong", "don_gia", "thanh_tien", "tong_cong",
            "tong_tien_hang", "thue_gtgt", "tong_thanh_toan", "so_tien_bang_chu", 
            "ghi_chu", "muc_dich_su_dung", "ton_kho", "nha_cung_cap", "muc_dich", "tr"
        ]
        
        for tpl in selected_templates:
            tpl_path = os.path.join(TEMPLATE_DIR, tpl)
            raw_tags = exhaustive_extract_tags(tpl_path) # Gọi hàm quét siêu mạnh
            clean_tags = set()
            for tag in raw_tags:
                t = str(tag).lower().strip()
                # Loại bỏ thẻ bảng kê và thẻ lập trình hệ thống
                if not ("%" in t or t.startswith("item.") or "endfor" in t or "for " in t or t in TABLE_TAGS):
                    clean_tags.add(t)
                    all_dynamic_tags.add(t)
            template_tags_map[tpl] = clean_tags
            
        if st.button("🚀 KÍCH HOẠT AI ĐỌC MẪU & ĐỐI CHIẾU", type="primary", use_container_width=True):
            with st.spinner("AI đang phân tích sâu dữ liệu hóa đơn..."):
                data = AIExtractor(API_KEY).extract_invoice_data(temp_input, expected_tags=all_dynamic_tags) 
                if "error" not in data: st.session_state['data'] = data
                else: st.error(data['error'])

        if 'data' in st.session_state:
            data = st.session_state['data']
            
            canh_bao = data.get("danh_sach_canh_bao", [])
            if canh_bao and len(canh_bao) > 0 and canh_bao[0].strip() != "":
                st.error("🚨 **CẢNH BÁO TỪ AI!**")
                for loi in canh_bao: st.warning(f"⚠️ {loi}")
            else:
                st.success("✅ AI báo cáo: Số liệu chuẩn khớp.")

            # Trải phẳng dữ liệu (Flatten) để đảm bảo móc được hết data dù AI gom vào nhóm nào
            ai_flat_data = {}
            if isinstance(data, dict):
                for k, v in data.items():
                    if isinstance(v, dict):
                        for sub_k, sub_v in v.items():
                            ai_flat_data[sub_k] = sub_v
                    elif not isinstance(v, list):
                        ai_flat_data[k] = v
            
            final_dynamic_data = {} 
            
            if template_tags_map:
                st.markdown("**📝 KIỂM DUYỆT TỪNG BIỂU MẪU:**")
                
                TAG_MAPPING = {
                    "ho_ten_nguoi_de_nghi": "Họ tên người đề nghị",
                    "bo_phan": "Bộ phận",
                    "bo_phan_cong_tac": "Bộ phận công tác",
                    "ly_do_de_nghi": "Lý do đề nghị",
                    "ly_do_thanh_toan": "Lý do thanh toán",
                    "hinh_thuc_thanh_toan": "Hình thức thanh toán/tạm ứng",
                    "thoi_han_hoan_ung": "Thời hạn hoàn ứng",
                    "ngay_lap_phieu": "Ngày lập phiếu",
                    "ngay_ky": "Ngày ký",
                    "thang_ky": "Tháng ký",
                    "nam_ky": "Năm ký",
                    "ngay_thang_nam": "Ngày tháng năm",
                    "so_phieu": "Số phiếu",
                    "so_hop_dong": "Số Hợp đồng",
                    "ten_cong_ty": "Tên công ty",
                    "ten_cong_ty_ben_b": "Tên công ty Bên B",
                    "nguoi_dai_dien_ben_b": "Người đại diện Bên B",
                    "chuc_vu_ben_b": "Chức vụ Bên B",
                    "dia_chi": "Địa chỉ",
                    "dia_chi_ben_b": "Địa chỉ Bên B",
                    "ma_so_thue": "Mã số thuế",
                    "mst_ben_b": "MST Bên B",
                    "dien_thoai": "Điện thoại",
                    "dien_thoai_ben_b": "Điện thoại Bên B",
                    "so_tai_khoan": "Số tài khoản",
                    "so_tai_khoan_ben_b": "Số tài khoản Bên B",
                    "ten_ngan_hang": "Tên ngân hàng",
                    "thoi_gian_thuc_hien": "Thời gian thực hiện",
                    "dia_diem_thuc_hien": "Địa điểm thực hiện",
                    "thoi_gian_bao_hanh": "Thời gian bảo hành",
                    "ty_le_tam_ung": "Tỷ lệ tạm ứng (%)",
                    "so_don_dat_hang": "Số đơn đặt hàng",
                    "so_de_xuat": "Số đề xuất",
                    "nguoi_phu_trach": "Người phụ trách",
                    "email_phu_trach": "Email phụ trách",
                    "nguoi_nhan_hang": "Người nhận hàng",
                    "sdt_nguoi_nhan": "SĐT người nhận",
                    "nguoi_de_xuat": "Người đề xuất"
                }
                
                # Biến lưu các tag đã được render thành công ở các file trước
                displayed_tags = set()
                
                for tpl in selected_templates:
                    tags_in_tpl = template_tags_map.get(tpl, set())
                    
                    with st.expander(f"📑 Biểu mẫu: {tpl}", expanded=True):
                        # Xử lý trường hợp file chỉ chứa bảng kê mà không có thông tin chung
                        if not tags_in_tpl:
                            st.info("✔️ Biểu mẫu này sử dụng hoàn toàn dữ liệu từ **Bảng Kê Hàng Hóa** ở bên dưới, không có trường thông tin động nào khác cần nhập.")
                        else:
                            cols = st.columns(3) 
                            for idx, tag in enumerate(sorted(tags_in_tpl)):
                                col = cols[idx % 3]
                                label = TAG_MAPPING.get(tag, tag.replace("_", " ").title())
                                
                                with col:
                                    if tag not in displayed_tags:
                                        # Lần đầu tiên tag này xuất hiện -> Cho phép nhập
                                        default_val = str(ai_flat_data.get(tag, "")) 
                                        final_dynamic_data[tag] = st.text_input(
                                            label, 
                                            value=default_val, 
                                            key=f"dyn_{tag}",
                                            placeholder=f"Nhập {label.lower()}..."
                                        )
                                        displayed_tags.add(tag)
                                    else:
                                        # Tag này đã xuất hiện ở file trước -> Làm mờ (Disabled) tránh người dùng nhập 2 lần
                                        current_val = st.session_state.get(f"dyn_{tag}", final_dynamic_data.get(tag, ""))
                                        st.text_input(
                                            f"{label} (Dùng chung)", 
                                            value=current_val, 
                                            key=f"readonly_{tpl}_{tag}",
                                            disabled=True,
                                            help="Trường thông tin này đã được bạn nhập ở biểu mẫu bên trên. Hệ thống sẽ tự động đồng bộ xuống file này."
                                        )
            
            st.markdown("**📦 Bảng Kê Hàng Hóa / Dịch Vụ:**")
            ds = data.get("danh_sach_hang_hoa", [])
            if not ds: ds = [{"ten_hang_hoa": ""}]
            
            fmt_data = []
            for item in ds:
                gia = float(item.get("don_gia") or 0)
                sl = float(item.get("so_luong") or 0)
                fmt_data.append({
                    "Tên hàng hóa": item.get("ten_hang_hoa", ""), 
                    "ĐVT": item.get("don_vi_tinh", ""), 
                    "Số lượng": sl, 
                    "Đơn giá": gia, 
                    "Thành tiền": item.get("thanh_tien", sl * gia)
                })
            
            edited_df = st.data_editor(pd.DataFrame(fmt_data), num_rows="dynamic", use_container_width=True, hide_index=True)
            danh_sach_da_chinh_sua = edited_df.to_dict('records')

            ds_chuan = []
            need_rerun = False
            
            for item in danh_sach_da_chinh_sua:
                ten_hh = str(item.get("Tên hàng hóa") or "").strip()
                if not ten_hh or ten_hh == "None": continue 
                
                sl = float(item.get("Số lượng") or 0)
                gia = float(item.get("Đơn giá") or 0)
                
                thanh_tien_chuan = sl * gia
                thanh_tien_hien_tai = float(item.get("Thành tiền") or 0)
                
                if thanh_tien_hien_tai != thanh_tien_chuan:
                    need_rerun = True

                ds_chuan.append({
                    "ten_hang_hoa": ten_hh, "don_vi_tinh": str(item.get("ĐVT") or ""),
                    "so_luong": sl, "don_gia": gia, "thanh_tien": thanh_tien_chuan, 
                    "muc_dich_su_dung": "", "ton_kho": "", "nha_cung_cap": "", "ghi_chu": "", "muc_dich": ""
                })

            if need_rerun:
                st.session_state['data']['danh_sach_hang_hoa'] = ds_chuan
                st.rerun()

            tong_tien_hang = 0.0
            for item in ds_chuan:
                val = item.get("thanh_tien", 0)
                if pd.notna(val):
                    try: tong_tien_hang += float(val)
                    except ValueError: pass
            
            st.markdown("---")
            col_t1, col_t2, col_t3 = st.columns(3)
            with col_t1: st.metric("Cộng tiền hàng:", f"{tong_tien_hang:,.0f} đ")
            with col_t2: 
                thue_suat = st.number_input("Thuế suất GTGT (%)", value=8)
                tien_thue = tong_tien_hang * (thue_suat / 100)
                st.metric("Tiền thuế GTGT:", f"{tien_thue:,.0f} đ")
            with col_t3:
                tong_thanh_toan = tong_tien_hang + tien_thue
                if pd.isna(tong_thanh_toan): tong_thanh_toan = 0.0
                st.metric("Tổng cộng thanh toán:", f"{tong_thanh_toan:,.0f} đ")
                chu_so_tien = doc_so_tien_vn(tong_thanh_toan)
                st.markdown(f"*(Bằng chữ: {chu_so_tien})*")

            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("✅ XUẤT BỘ HỒ SƠ CHUẨN ERP", type="primary", use_container_width=True):
                
                for item in ds_chuan:
                    tong_cong_mon = item["thanh_tien"] * (1 + (thue_suat / 100))
                    item["tong_cong"] = f"{tong_cong_mon:,.0f}"
                
                final_data = {
                    **final_dynamic_data, 
                    "tong_tien_hang": tong_tien_hang, "thue_gtgt": tien_thue, 
                    "tong_thanh_toan": tong_thanh_toan, "tong_cong": tong_thanh_toan,
                    "so_tien_bang_chu": chu_so_tien, 
                    "danh_sach_hang_hoa": ds_chuan
                }
                
                try:
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                        for tpl in selected_templates:
                            tpl_path = os.path.join(TEMPLATE_DIR, tpl)
                            out_file = f"BaoCao_{tpl}"
                            
                            if tpl.endswith('.docx'): DocumentBuilder.fill_word_template(tpl_path, out_file, final_data)
                            else: DocumentBuilder.fill_excel_template(tpl_path, out_file, final_data)
                            
                            zip_file.write(out_file, arcname=out_file)
                            os.remove(out_file) 
                    
                    st.success("Tuyệt vời! Toàn bộ file đã được xử lý xong.")
                    st.download_button(
                        label="💾 BẤM VÀO ĐÂY ĐỂ TẢI BỘ HỒ SƠ (FILE ZIP)", 
                        data=zip_buffer.getvalue(), file_name="Bo_Ho_So_Ketoan.zip", 
                        mime="application/zip", type="secondary", use_container_width=True
                    )
                    
                    ten_doi_tac = final_dynamic_data.get('ten_cong_ty_ben_b', ai_flat_data.get('ten_cong_ty', ''))
                    display_zalo_message(ten_doi_tac, tong_thanh_toan)

                except Exception as e:
                    st.error(f"Lỗi: {str(e)}")
