from docx import Document
import openpyxl
from copy import copy, deepcopy
import re

class DocumentBuilder:
    @staticmethod
    def fill_word_template(template_path, output_path, data_dict):
        doc = Document(template_path)
        danh_sach = data_dict.get("danh_sach_hang_hoa", [])

        if danh_sach:
            for table in doc.tables:
                template_row = next((r for r in table.rows if "{ten_hang_hoa}" in "".join(c.text for c in r.cells) or "{stt}" in "".join(c.text for c in r.cells)), None)
                if template_row:
                    parent_tbl = template_row._tr.getparent()
                    for idx, item in enumerate(danh_sach):
                        new_tr = deepcopy(template_row._tr)
                        parent_tbl.insert(parent_tbl.index(template_row._tr), new_tr)
                        from docx.table import _Row
                        new_row = _Row(new_tr, table)
                        
                        for cell in new_row.cells:
                            for p in cell.paragraphs:
                                if "{stt}" in p.text: p.text = p.text.replace("{stt}", str(idx+1))
                                for k, v in item.items():
                                    search_key = f"{{{k}}}"
                                    if search_key in p.text:
                                        val_str = f"{v:,.0f}" if isinstance(v, (int, float)) and k in ["so_luong", "don_gia", "thanh_tien", "tong_cong"] else str(v if v is not None else "")
                                        p.text = p.text.replace(search_key, val_str)
                                p.text = re.sub(r'\{[a-z_]+\}', '', p.text)
                    parent_tbl.remove(template_row._tr)

        def replace_text(paragraph, key, value):
            search_key = f"{{{key}}}"
            if search_key in paragraph.text:
                val_str = f"{value:,.0f}" if isinstance(value, (int, float)) and key in ["tong_tien_hang", "thue_gtgt", "tong_thanh_toan", "tong_cong"] else str(int(value) if isinstance(value, float) and value.is_integer() else (value or ""))
                paragraph.text = paragraph.text.replace(search_key, val_str)

        for key_name, value in data_dict.items():
            if isinstance(value, list): continue
            for p in doc.paragraphs: replace_text(p, key_name, value)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs: replace_text(p, key_name, value)
        
        doc.save(output_path)
        return True

    @staticmethod
    def fill_excel_template(template_path, output_path, data_dict):
        wb = openpyxl.load_workbook(template_path)
        sheet = wb.active
        danh_sach = data_dict.get("danh_sach_hang_hoa", [])
        table_start_row = None
        col_map = {}
        list_keys = ["stt", "ten_hang_hoa", "muc_dich", "ton_kho", "nha_cung_cap", "don_vi_tinh", "so_luong", "don_gia", "gia_niem_yet", "gia_mua", "thanh_tien", "ghi_chu", "tong_cong"]
        
        for r in range(1, 101):
            for c in range(1, 51):
                try:
                    val = sheet.cell(row=r, column=c).value
                    if val and isinstance(val, str):
                        for key in list_keys:
                            if f"{{{key}}}" in val:
                                col_map[key] = c
                                table_start_row = r
                except Exception: pass
            if table_start_row: break

        merged_bounds = [m.bounds for m in sheet.merged_cells.ranges]
        merged_coords = [m.coord for m in list(sheet.merged_cells.ranges)]
        for coord in merged_coords: sheet.unmerge_cells(coord)
        
        so_dong_them = len(danh_sach) - 1 if danh_sach else 0

        if table_start_row and so_dong_them > 0:
            sheet.insert_rows(table_start_row + 1, so_dong_them)
            for i in range(1, so_dong_them + 1):
                for c in range(1, sheet.max_column + 1):
                    src_cell, tgt_cell = sheet.cell(row=table_start_row, column=c), sheet.cell(row=table_start_row + i, column=c)
                    if src_cell.has_style:
                        tgt_cell.font, tgt_cell.border, tgt_cell.fill, tgt_cell.number_format, tgt_cell.alignment = copy(src_cell.font), copy(src_cell.border), copy(src_cell.fill), copy(src_cell.number_format), copy(src_cell.alignment)

        for min_col, min_row, max_col, max_row in merged_bounds:
            if table_start_row and min_row > table_start_row:
                sheet.merge_cells(start_row=min_row + so_dong_them, start_column=min_col, end_row=max_row + so_dong_them, end_column=max_col)
            elif table_start_row and min_row == table_start_row:
                for i in range(so_dong_them + 1): sheet.merge_cells(start_row=min_row + i, start_column=min_col, end_row=max_row + i, end_column=max_col)
            else:
                sheet.merge_cells(start_row=min_row, start_column=min_col, end_row=max_row, end_column=max_col)

        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    for key, value in data_dict.items():
                        if isinstance(value, list): continue
                        if f"{{{key}}}" in cell.value:
                            val_str = f"{value:,.0f}" if isinstance(value, (int, float)) and key in ["tong_tien_hang", "thue_gtgt", "tong_thanh_toan", "tong_cong"] else str(int(value) if isinstance(value, float) and value.is_integer() else (value or ""))
                            cell.value = cell.value.replace(f"{{{key}}}", val_str)

        if table_start_row and col_map:
            for idx, item in enumerate(danh_sach):
                for key_name, col_idx in col_map.items():
                    val = idx + 1 if key_name == "stt" else (float(item.get(key_name, 0) or 0) if key_name in ["so_luong", "don_gia", "thanh_tien", "tong_cong"] else item.get(key_name, ""))
                    c = sheet.cell(row=table_start_row + idx, column=col_idx)
                    c.value = val
                    if key_name in ["so_luong", "don_gia", "thanh_tien", "tong_cong"]: c.number_format = '#,##0'
                    
        wb.save(output_path)
        return True