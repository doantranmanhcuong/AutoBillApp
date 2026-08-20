import re
from docx import Document
import openpyxl

def doc_so_tien_vn(n):
    if n == 0: return "Không đồng"
    units = ["", " nghìn", " triệu", " tỷ", " nghìn tỷ", " triệu tỷ"]
    words = ["không", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]

    def doc_3_so(num, read_zero_hundred=False):
        h = num // 100
        t = (num % 100) // 10
        u = num % 10
        res = ""
        if h > 0 or read_zero_hundred: res += words[h] + " trăm "
        if t > 1:
            res += words[t] + " mươi "
            if u == 1: res += "mốt "
            elif u == 5: res += "lăm "
            elif u > 0: res += words[u] + " "
        elif t == 1:
            res += "mười "
            if u == 5: res += "lăm "
            elif u > 0: res += words[u] + " "
        elif t == 0 and u > 0 and (h > 0 or read_zero_hundred): res += "lẻ " + words[u] + " "
        elif t == 0 and u > 0: res += words[u] + " "
        return res.strip()

    s = ""
    group = 0
    n_int = int(n)
    while n_int > 0:
        chunk = n_int % 1000
        n_int = n_int // 1000
        if chunk > 0:
            chunk_str = doc_3_so(chunk, read_zero_hundred=(n_int > 0))
            s = chunk_str + units[group] + " " + s
        group += 1
    
    return s.strip().capitalize() + " đồng chẵn."

def get_tags_from_template(file_path):
    tags = set()
    pattern = r'\{([^{}]+)\}' 
    try:
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            for p in doc.paragraphs: tags.update(re.findall(pattern, p.text))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs: tags.update(re.findall(pattern, p.text))
        elif file_path.endswith('.xlsx'):
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            tags.update(re.findall(pattern, cell.value))
    except Exception: pass

    ignore_tags = {"stt", "ten_hang_hoa", "muc_dich", "ton_kho", "nha_cung_cap", 
                   "don_vi_tinh", "so_luong", "don_gia", "gia_niem_yet", "gia_mua", 
                   "thanh_tien", "ghi_chu", "tong_tien_hang", "thue_gtgt", 
                   "tong_thanh_toan", "tong_cong", "so_tien_bang_chu"} 
    return [t for t in tags if t not in ignore_tags]